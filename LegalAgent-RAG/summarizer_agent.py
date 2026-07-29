import os
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from unstructured.partition.pdf import partition_pdf

class MixedContentLoader:
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.ext = os.path.splitext(file_path)[1].lower()

    def load(self):
        if self.ext == ".pdf":
            return self._load_pdf()
        elif self.ext == ".docx":
            return self._load_docx()
        elif self.ext == ".txt":
            return self._load_txt()
        else:
            raise ValueError(f"Unsupported file type: {self.ext}")

    def _load_pdf(self):
        print(" Trying PDF parsing without OCR (strategy='hi_res')...")

        elements = partition_pdf(
            filename=self.file_path,
            strategy="hi_res",            # No OCR here
            extract_images_in_pdf=False
        )

        chunks = self._elements_to_chunks(elements)

        # # Fallback to OCR if very little content was found
        # if len(chunks) < 5:
        #     print(" Not enough content found — retrying with OCR (strategy='ocr_only')...")
        #     try:
        #         elements = partition_pdf(
        #             filename=self.file_path,
        #             strategy="ocr_only",     # OCR enabled here
        #             ocr_languages="eng",
        #             extract_images_in_pdf=False
        #         )
        #         chunks = self._elements_to_chunks(elements)
        #         print(" OCR fallback succeeded.")
        #     except Exception as e:
        #         print(f" OCR fallback failed: {e}")

        return chunks, "pdf"

    def _elements_to_chunks(self, elements):
        chunks = []
        for el in elements:
            # Your previous code handled tables differently, so let's keep that logic:
            text = str(el).strip() if not hasattr(el, "text") else el.text.strip()
            if not text:
                continue

            # Detect and process tables if possible
            # unstructured returns element.category or element.type differently depending on version
            el_type = getattr(el, "category", None) or getattr(el, "type", "").lower()

            if el_type == "table" or el_type == "tableelement":
                # Defensive fallback if table_rows and cells exist
                try:
                    table_text = ""
                    for row in el.table_rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells)
                        table_text += row_text + "\n"
                    chunks.append({
                        "type": "table",
                        "content": table_text.strip()
                    })
                except Exception:
                    # fallback to raw text if no structured table data
                    chunks.append({
                        "type": "table",
                        "content": text
                    })
            else:
                chunks.append({
                    "type": "text",
                    "content": text
                })

        return chunks

    def _load_docx(self):
        chunks = []
        doc = DocxDocument(self.file_path)
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                chunks.append({
                    "type": "text",
                    "content": text
                })

        # Tables
        for table in doc.tables:
            table_text = ""
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_text += " | ".join(row_data) + "\n"
            if table_text.strip():
                chunks.append({
                    "type": "table",
                    "content": table_text.strip()
                })

        return chunks, "docx"

    def _load_txt(self):
        chunks = []
        with open(self.file_path, "r", encoding="utf-8") as f:
            text = f.read().strip()
            if text:
                chunks.append({
                    "type": "text",
                    "content": text
                })
        return chunks, "txt"
